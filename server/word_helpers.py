from io import BytesIO
from typing import List
from docx import Document
from docx.oxml.ns import qn
from .markdown_helper import wrap_text

def extract_text_bytes(data: bytes, text_inject_max_chars: int) -> tuple[str, bool]:
    # Decode bytes as text for injection. Returns (text, was_truncated).
    if data.count(b'\x00') > max(8, len(data) // 100):
        raise ValueError('binary-like payload (NUL bytes)')

    try:
        txt = data.decode('utf-8')
    except UnicodeDecodeError:
        txt = data.decode('utf-8', errors='replace')

    truncated = False
    if len(txt) > text_inject_max_chars:
        txt = txt[:text_inject_max_chars] + f"\n\n[...truncated; exceeded text_inject_max_chars={text_inject_max_chars}]"
        truncated = True
    return txt, truncated

def _xpath(el, expr: str):
    """
    python-docx oxml elements usually provide .xpath(expr) with namespaces prewired.
    In case we're on raw lxml, fall back to namespaces kw.
    """
    try:
        return el.xpath(expr)
    except TypeError:
        # raw lxml path (rare in python-docx, but safe)
        return el.xpath(expr, namespaces=WORDX_NS)

# Namespaces used by WordprocessingML
WORDX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def extract_docx_markdown(data: bytes, text_inject_max_chars: int) -> str:
    """
    Best-effort DOCX -> markdown-ish text + JSON-ish table blocks + image annotations.
    Improvements vs prior:
      - Preserves Word hyperlinks as [text](url) by resolving w:hyperlink r:id
      - Uses your house dialect via wrap_text() (bold=**, underline=__ etc.)
    """
    if Document is None:
        raise RuntimeError("python-docx is not installed (Document import failed)")

    import json as _json
    import re as _re

    doc = Document(BytesIO(data))
    out_lines: List[str] = []

    def _style_prefix(paragraph) -> str:
        style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
        if style_name.startswith("heading"):
            m = _re.search(r"(\d+)", style_name)
            level = int(m.group(1)) if m else 1
            level = max(1, min(level, 6))
            return ("#" * level) + " "
        return ""

    def _build_run_map(paragraph) -> dict[int, object]:
        # Map underlying XML run element identity -> python-docx Run object
        m = {}
        for run in paragraph.runs:
            try:
                m[id(run._r)] = run
            except Exception:
                pass
        return m

    def _render_run(run_obj) -> str:
        t = getattr(run_obj, "text", "") or ""
        if not t:
            return ""
        return wrap_text(
            t,
            bool(getattr(run_obj, "bold", False)),
            bool(getattr(run_obj, "italic", False)),
            bool(getattr(run_obj, "underline", False)),
            False,  # strike (TODO)
            False,  # spoiler (unused)
        )

    def _resolve_hyperlink_url(rel_id: str) -> str:
        if not rel_id:
            return ""
        try:
            rel = doc.part.rels.get(rel_id)
            if rel is None:
                return ""
            # target_ref is typically a URL for external links
            return getattr(rel, "target_ref", "") or ""
        except Exception:
            return ""

    def _render_paragraph_with_links(paragraph) -> str:
        prefix = _style_prefix(paragraph)
        run_map = _build_run_map(paragraph)

        parts: List[str] = []

        # Walk top-level children: w:r and w:hyperlink
        try:
            children = paragraph._p.xpath("./w:r | ./w:hyperlink", namespaces=WORDX_NS)
        except Exception:
            # Fallback: no XML access, revert to old behavior
            children = []
        if not children:
            # Old behavior fallback: concatenate paragraph.runs
            for r in paragraph.runs:
                s = _render_run(r)
                if s:
                    parts.append(s)
            return (prefix + "".join(parts)).rstrip()

        for el in children:
            tag = getattr(el, "tag", "")
            if tag.endswith("}r"):
                run_obj = run_map.get(id(el))
                if run_obj is not None:
                    s = _render_run(run_obj)
                    if s:
                        parts.append(s)
                else:
                    # fallback: try to read text nodes directly
                    try:
                        texts = el.xpath(".//w:t/text()", namespaces=WORDX_NS)
                        raw = "".join(texts)
                        if raw:
                            parts.append(raw)
                    except Exception:
                        pass

            elif tag.endswith("}hyperlink"):
                # Resolve URL
                rid = el.get(f"{{{WORDX_NS['r']}}}id") or ""
                url = _resolve_hyperlink_url(rid)

                # Render runs inside hyperlink
                link_parts: List[str] = []
                try:
                    run_elems = el.xpath(".//w:r", namespaces=WORDX_NS)
                except Exception:
                    run_elems = []

                for r_el in run_elems:
                    run_obj = run_map.get(id(r_el))
                    if run_obj is not None:
                        s = _render_run(run_obj)
                        if s:
                            link_parts.append(s)
                    else:
                        try:
                            texts = r_el.xpath(".//w:t/text()", namespaces=WORDX_NS)
                            raw = "".join(texts)
                            if raw:
                                link_parts.append(raw)
                        except Exception:
                            pass

                visible = "".join(link_parts).strip()
                if not visible:
                    continue

                if url:
                    # Markdown link. Keep the visible text as-is (already contains styling markers).
                    parts.append(f"[{visible}]({url})")
                else:
                    # If we can’t resolve target, at least keep the visible text.
                    parts.append(visible)

        return (prefix + "".join(parts)).rstrip()

    for p in doc.paragraphs:
        line = _render_paragraph_with_links(p).strip()

        # Images in this paragraph (best-effort) — keep your existing marker format
        try:
            drawings = _xpath(p._p, ".//w:drawing")
            if drawings:
                for d in drawings:
                    alt = ""
                    name = ""
                    try:
                        docPr = d.xpath(".//wp:docPr", namespaces=WORDX_NS)
                        if docPr:
                            alt = docPr[0].get("descr") or docPr[0].get("title") or ""
                            name = docPr[0].get("name") or ""
                    except Exception:
                        pass

                    rel_id = ""
                    try:
                        blips = d.xpath(".//a:blip", namespaces=WORDX_NS)
                        if blips:
                            rel_id = blips[0].get(f"{{{WORDX_NS['r']}}}embed") or ""
                    except Exception:
                        pass

                    img_name = (name or "").strip()
                    if rel_id:
                        try:
                            part = doc.part.related_parts[rel_id]
                            img_name = str(getattr(part, "partname", img_name)).split("/")[-1] or img_name
                        except Exception:
                            pass
                    if not img_name:
                        img_name = "image"

                    marker = "{ image: " + _json.dumps(img_name) + ", alt: " + _json.dumps(alt) + " }"
                    line = (line + " " + marker).strip() if line else marker
        except Exception:
            pass

        if line:
            out_lines.append(line)

    # Tables (same behavior as your current implementation)
    table_blocks: List[str] = []
    try:
        for ti, table in enumerate(doc.tables, start=1):
            rows = [[(cell.text or "").strip() for cell in row.cells] for row in table.rows]
            headers = rows[0] if rows else []
            header_ok = bool(headers) and all(h.strip() for h in headers) and (len(set(headers)) == len(headers))

            if header_ok and len(rows) > 1:
                objs = []
                for r in rows[1:]:
                    objs.append({headers[i]: (r[i] if i < len(r) else "") for i in range(len(headers))})
                payload = {f"table_{ti}": objs}
            else:
                payload = {f"table_{ti}": rows}

            table_blocks.append(_json.dumps(payload, ensure_ascii=False, indent=2))
    except Exception:
        pass

    if table_blocks:
        out_lines.append("")
        out_lines.append("Attachment tables (JSON-ish):")
        out_lines.extend(table_blocks)

    txt = "\n".join(out_lines).strip()
    if len(txt) > text_inject_max_chars:
        txt = txt[:text_inject_max_chars] + (
            f"\n\n[...truncated; exceeded text_inject_max_chars={text_inject_max_chars}]"
        )
    return txt

if (False): # Legacy replaced by above
    def extract_docx_markdown(data: bytes, text_inject_max_chars: int) -> str:
        # Best-effort DOCX -> markdown-ish text + JSON-ish table blocks + image annotations.
        if Document is None:
            raise RuntimeError('python-docx is not installed (Document import failed)')

        import json as _json
        import re as _re

        doc = Document(BytesIO(data))
        out_lines: List[str] = []

        for p in doc.paragraphs:
            style_name = (getattr(getattr(p, 'style', None), 'name', '') or '').lower()
            prefix = ''
            if style_name.startswith('heading'):
                m = _re.search(r'(\d+)', style_name)
                level = int(m.group(1)) if m else 1
                level = max(1, min(level, 6))
                prefix = '#' * level + ' '

            parts: List[str] = []
            for r in p.runs:
                t = r.text or ''
                if not t:
                    continue
                # TODO support strikethrough; don't think Word has spoiler text
                parts.append(wrap_text(t, bool(getattr(r, 'bold', False)), bool(getattr(r, 'italic', False)), bool(getattr(r, 'underline', False)), False, False))
            line = (prefix + ''.join(parts)).rstrip()

            # Images in this paragraph (best-effort)
            try:
                drawings = p._p.xpath('.//w:drawing')
                if drawings:
                    for d in drawings:
                        alt = ''
                        name = ''
                        try:
                            docPr = d.xpath('.//wp:docPr')
                            if docPr:
                                alt = docPr[0].get('descr') or docPr[0].get('title') or ''
                                name = docPr[0].get('name') or ''
                        except Exception:
                            pass

                        rel_id = ''
                        try:
                            blips = d.xpath('.//a:blip')
                            if blips:
                                rel_id = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed') or ''
                        except Exception:
                            pass

                        img_name = (name or '').strip()
                        if rel_id:
                            try:
                                part = doc.part.related_parts[rel_id]
                                img_name = str(getattr(part, 'partname', img_name)).split('/')[-1] or img_name
                            except Exception:
                                pass
                        if not img_name:
                            img_name = 'image'

                        marker = '{ image: ' + _json.dumps(img_name) + ', alt: ' + _json.dumps(alt) + ' }'
                        line = (line + ' ' + marker).strip() if line else marker
            except Exception:
                pass

            if line:
                out_lines.append(line)

        # Tables
        table_blocks: List[str] = []
        try:
            for ti, table in enumerate(doc.tables, start=1):
                rows = [[(cell.text or '').strip() for cell in row.cells] for row in table.rows]
                headers = rows[0] if rows else []
                header_ok = bool(headers) and all(h.strip() for h in headers) and (len(set(headers)) == len(headers))

                if header_ok and len(rows) > 1:
                    objs = []
                    for r in rows[1:]:
                        objs.append({headers[i]: (r[i] if i < len(r) else '') for i in range(len(headers))})
                    payload = {f'table_{ti}': objs}
                else:
                    payload = {f'table_{ti}': rows}

                table_blocks.append(_json.dumps(payload, ensure_ascii=False, indent=2))
        except Exception:
            pass

        if table_blocks:
            out_lines.append('')
            out_lines.append('Attachment tables (JSON-ish):')
            out_lines.extend(table_blocks)

        txt = '\n'.join(out_lines).strip()
        if len(txt) > text_inject_max_chars:
            txt = txt[:text_inject_max_chars] + f"\n\n[...truncated; exceeded text_inject_max_chars={text_inject_max_chars}]"
        return txt
